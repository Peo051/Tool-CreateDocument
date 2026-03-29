#!/usr/bin/env python3
"""
CÔNG CỤ TỰ ĐỘNG TẠO BÁO CÁO PRO - FILE CHÍNH
Tích hợp đầy đủ: Nội dung + Biểu đồ + Export PDF
"""

import os
import sys
import json
import argparse
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Import modules
try:
    from algorithm_content import ALGORITHMS, get_algorithm_content
    from diagram_generator import DiagramGenerator
except ImportError:
    print("⚠️  Không tìm thấy module phụ trợ!")
    print("💡 Đảm bảo các file sau cùng thư mục:")
    print("   - algorithm_content.py")
    print("   - diagram_generator.py")
    sys.exit(1)


class ProReportGenerator:
    """Generator báo cáo PRO với đầy đủ tính năng"""
    
    def __init__(self, config_file='report_config.json'):
        print("🚀 Khởi tạo Pro Report Generator...")
        self.config = self.load_config(config_file)
        self.doc = Document()
        self.diagram_gen = DiagramGenerator()
        self.setup_document()
        
        # Tạo thư mục diagrams
        os.makedirs('diagrams', exist_ok=True)
        os.makedirs('output', exist_ok=True)
        
    def load_config(self, config_file):
        """Load cấu hình"""
        if os.path.exists(config_file):
            with open(config_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        
        print("⚠️  Không tìm thấy report_config.json, dùng config mặc định")
        return self.get_default_config()
    
    def get_default_config(self):
        """Config mặc định"""
        return {
            "title": "Maze Duel - BO3 / Gemini_Game",
            "subtitle": "Ứng dụng thuật toán AI trong game mê cung đối kháng",
            "students": [
                {"name": "Trần Dương Gia Bảo", "id": "2001240039"},
                {"name": "Nguyễn Thế Anh", "id": "2001240015"}
            ],
            "advisor": "TS. Trần Việt Hùng",
            "class": "DHCN01",
            "faculty": "Công nghệ thông tin",
            "university": "Đại học Công Thương TP. HCM",
            "github_repo": "https://github.com/GiaBao051/Gemini_Game",
            "demo_url": "https://giabao051.github.io/Gemini_Game/",
            "algorithms": ["DFS", "BFS", "A*", "Dijkstra", "CSP"],
            "technologies": ["HTML5", "JavaScript", "Canvas API", "MQTT"]
        }
    
    def setup_document(self):
        """Thiết lập document"""
        sections = self.doc.sections
        for section in sections:
            section.page_height = Cm(29.7)
            section.page_width = Cm(21)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2)
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
        
        self.setup_styles()
    
    def setup_styles(self):
        """Thiết lập styles"""
        styles = self.doc.styles
        
        normal = styles['Normal']
        normal.font.name = 'Times New Roman'
        normal.font.size = Pt(13)
        normal.paragraph_format.line_spacing = 1.5
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        for i in range(1, 4):
            h = styles[f'Heading {i}']
            h.font.name = 'Times New Roman'
            h.font.bold = True
            if i == 1:
                h.font.size = Pt(16)
                h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif i == 2:
                h.font.size = Pt(14)
            else:
                h.font.size = Pt(13)
    
    def add_centered_text(self, text, size, bold=False):
        """Thêm text căn giữa"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
    
    def add_cover_page(self):
        """Trang bìa"""
        print("📄 Tạo trang bìa...")
        
        self.add_centered_text('BỘ CÔNG THƯƠNG', 13, True)
        self.add_centered_text(f'TRƯỜNG {self.config["university"].upper()}', 13, True)
        self.add_centered_text(f'KHOA {self.config["faculty"].upper()}', 13, True)
        self.add_centered_text('─' * 30, 13)
        
        for _ in range(3):
            self.doc.add_paragraph()
        
        self.add_centered_text('BÁO CÁO ĐỒ ÁN MÔN HỌC', 16, True)
        self.add_centered_text('TRÍ TUỆ NHÂN TẠO', 16, True)
        self.doc.add_paragraph()
        
        self.add_centered_text(f'ĐỀ TÀI: {self.config["title"]}', 14, True)
        self.add_centered_text(self.config["subtitle"], 13)
        
        for _ in range(3):
            self.doc.add_paragraph()
        
        self.add_centered_text(f'Giảng viên hướng dẫn: {self.config["advisor"]}', 13)
        self.doc.add_paragraph()
        
        self.add_centered_text('Sinh viên thực hiện:', 13, True)
        for student in self.config["students"]:
            self.add_centered_text(f'{student["name"]} - MSSV: {student["id"]}', 13)
        
        for _ in range(2):
            self.doc.add_paragraph()
        
        self.add_centered_text(f'Lớp: {self.config["class"]}', 13)
        
        for _ in range(2):
            self.doc.add_paragraph()
        
        now = datetime.now()
        self.add_centered_text(f'TP. HỒ CHÍ MINH, tháng {now.month} năm {now.year}', 13)
        
        self.doc.add_page_break()
    
    def add_acknowledgment(self):
        """Lời cảm ơn"""
        print("💝 Thêm lời cảm ơn...")
        
        self.doc.add_heading('LỜI CẢM ƠN', 1)
        
        text = f"""Nhóm chúng em xin chân thành cảm ơn quý Thầy/Cô Khoa {self.config['faculty']}, đặc biệt là {self.config['advisor']} đã tận tình hướng dẫn, giúp đỡ nhóm trong suốt quá trình thực hiện đồ án môn Trí tuệ nhân tạo.

Qua đồ án này, nhóm đã có cơ hội áp dụng các kiến thức lý thuyết vào thực tiễn, hiểu sâu hơn về các thuật toán AI và cách chúng hoạt động trong một hệ thống thực tế.

Mặc dù đã cố gắng hết sức, báo cáo vẫn không tránh khỏi những thiếu sót. Nhóm rất mong nhận được sự góp ý của quý Thầy/Cô để hoàn thiện hơn.

Nhóm xin chân thành cảm ơn!"""
        
        self.doc.add_paragraph(text)
        self.doc.add_page_break()
    
    def add_commitment(self):
        """Lời cam đoan"""
        print("✍️  Thêm lời cam đoan...")
        
        self.doc.add_heading('LỜI CAM ĐOAN', 1)
        
        text = """Nhóm chúng em xin cam đoan đây là sản phẩm nghiên cứu của riêng nhóm. Các số liệu, kết quả nêu trong báo cáo là trung thực và chưa từng được ai công bố trong bất kỳ công trình nào khác.

Nhóm xin hoàn toàn chịu trách nhiệm về lời cam đoan này."""
        
        self.doc.add_paragraph(text)
        self.doc.add_paragraph()
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run('Sinh viên thực hiện')
        run.bold = True
        
        for student in self.config["students"]:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(student["name"])
        
        self.doc.add_page_break()
    
    def add_table_of_contents(self):
        """Mục lục"""
        print("📑 Tạo mục lục...")
        
        self.doc.add_heading('MỤC LỤC', 1)
        
        toc_items = [
            "PHẦN MỞ ĐẦU",
            "CHƯƠNG 1: TỔNG QUAN",
            "CHƯƠNG 2: THUẬT TOÁN TÌM KIẾM",
            "CHƯƠNG 3: BÀI TOÁN RÀNG BUỘC",
            "CHƯƠNG 4: XỬ LÝ TRI THỨC",
            "CHƯƠNG 5: TRIỂN KHAI",
            "KẾT LUẬN",
            "TÀI LIỆU THAM KHẢO"
        ]
        
        for item in toc_items:
            p = self.doc.add_paragraph()
            run = p.add_run(item)
            run.bold = True
        
        self.doc.add_page_break()
    
    def add_introduction(self):
        """Phần mở đầu"""
        print("📝 Viết phần mở đầu...")
        
        self.doc.add_heading('PHẦN MỞ ĐẦU', 1)
        
        self.doc.add_heading('1. Lý do chọn đề tài', 2)
        intro = f"""Trong bối cảnh công nghệ phát triển mạnh mẽ, Trí tuệ nhân tạo (AI) đã và đang trở thành một phần không thể thiếu trong nhiều lĩnh vực, đặc biệt là trong ngành công nghiệp game.

Đề tài "{self.config['title']}" được nhóm lựa chọn xuất phát từ mong muốn tìm hiểu sâu về cách các thuật toán AI hoạt động trong một hệ thống thực tế. Thay vì chỉ nghiên cứu lý thuyết, nhóm muốn xây dựng một ứng dụng game hoàn chỉnh để minh họa cách các thuật toán phối hợp với nhau."""
        
        self.doc.add_paragraph(intro)
        
        self.doc.add_heading('2. Mục tiêu nghiên cứu', 2)
        objectives = [
            "Nghiên cứu và triển khai các thuật toán AI cơ bản",
            "Áp dụng thuật toán CSP và Graph Coloring",
            "Xây dựng hệ thống AI đa tầng",
            "Triển khai chế độ online",
            "Tạo công cụ debug và visualization"
        ]
        for obj in objectives:
            self.doc.add_paragraph(obj, style='List Bullet')
        
        self.doc.add_page_break()
    
    def add_algorithm_chapter(self, algo_name, chapter_num, section_num):
        """Thêm phần phân tích thuật toán"""
        print(f"   📊 Phân tích {algo_name}...")
        
        content = get_algorithm_content(algo_name)
        if not content:
            print(f"   ⚠️  Không tìm thấy nội dung cho {algo_name}")
            return
        
        self.doc.add_heading(f'{chapter_num}.{section_num}. {content["title"]}', 2)
        
        # Mục tiêu
        self.doc.add_heading(f'{chapter_num}.{section_num}.1. Mục tiêu', 3)
        self.doc.add_paragraph(content["purpose"])
        
        # Lý thuyết
        self.doc.add_heading(f'{chapter_num}.{section_num}.2. Cơ sở lý thuyết', 3)
        self.doc.add_paragraph(content["theory"])
        
        # Độ phức tạp
        self.doc.add_heading(f'{chapter_num}.{section_num}.3. Độ phức tạp', 3)
        self.doc.add_paragraph(content["complexity"])
        
        # Giả mã
        self.doc.add_heading(f'{chapter_num}.{section_num}.4. Giả mã', 3)
        self.add_code_block(content["pseudocode"])
        
        # Ưu điểm
        self.doc.add_heading(f'{chapter_num}.{section_num}.5. Ưu điểm', 3)
        for adv in content["advantages"]:
            self.doc.add_paragraph(adv, style='List Bullet')
        
        # Hạn chế
        self.doc.add_heading(f'{chapter_num}.{section_num}.6. Hạn chế', 3)
        for lim in content["limitations"]:
            self.doc.add_paragraph(lim, style='List Bullet')
        
        # Thêm biểu đồ nếu có
        self.add_diagram_for_algorithm(algo_name)
    
    def add_code_block(self, code):
        """Thêm code block"""
        p = self.doc.add_paragraph()
        run = p.add_run(code)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F0F0F0')
        p._element.get_or_add_pPr().append(shading)
        
        p.paragraph_format.left_indent = Inches(0.5)
    
    def add_diagram_for_algorithm(self, algo_name):
        """Thêm biểu đồ cho thuật toán"""
        diagram_map = {
            'DFS': 'create_maze_example',
            'BFS': 'create_bfs_visualization',
            'A*': 'create_astar_comparison',
            'Dijkstra': 'create_dijkstra_cost_map',
            'CSP': 'create_csp_example'
        }
        
        if algo_name in diagram_map:
            try:
                method = getattr(self.diagram_gen, diagram_map[algo_name])
                img_buffer = method()
                
                # Lưu file với tên an toàn
                safe_name = algo_name.lower().replace('*', 'star').replace('/', '_')
                img_path = f'diagrams/{safe_name}_diagram.png'
                with open(img_path, 'wb') as f:
                    f.write(img_buffer.read())
                
                # Thêm vào document
                self.doc.add_paragraph()
                self.doc.add_picture(img_path, width=Inches(5.5))
                
                # Caption
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f'Hình: Minh họa thuật toán {algo_name}')
                run.italic = True
                run.font.size = Pt(11)
                
                print(f"      ✅ Đã thêm biểu đồ {algo_name}")
            except Exception as e:
                print(f"      ⚠️  Lỗi tạo biểu đồ {algo_name}: {e}")
    
    def add_conclusion(self):
        """Kết luận"""
        print("🎯 Viết kết luận...")
        
        self.doc.add_heading('KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN', 1)
        
        conclusion = f"""Qua quá trình nghiên cứu và triển khai đồ án "{self.config['title']}", nhóm đã đạt được các kết quả quan trọng trong việc áp dụng các thuật toán AI vào một hệ thống game hoàn chỉnh.

Đồ án đã thành công trong việc tích hợp nhiều thuật toán AI khác nhau, từ sinh môi trường (DFS, CSP) đến tìm đường (BFS, A*, Dijkstra) và ra quyết định (Minimax, Utility Scoring).

Hướng phát triển tiếp theo bao gồm: tối ưu hiệu năng, thêm machine learning, phát triển mobile app và tournament mode."""
        
        self.doc.add_paragraph(conclusion)
        self.doc.add_page_break()
    
    def add_references(self):
        """Tài liệu tham khảo"""
        print("📖 Thêm tài liệu tham khảo...")
        
        self.doc.add_heading('TÀI LIỆU THAM KHẢO', 1)
        
        refs = [
            f"[1] Repository: {self.config['github_repo']}",
            f"[2] Demo: {self.config['demo_url']}",
            "[3] Stuart Russell, Peter Norvig. Artificial Intelligence: A Modern Approach. 4th Edition, 2020.",
            "[4] Thomas H. Cormen et al. Introduction to Algorithms. 3rd Edition, MIT Press, 2009."
        ]
        
        for ref in refs:
            self.doc.add_paragraph(ref, style='List Number')
    
    def generate_full_report(self, output_file='output/Bao_Cao_Full_Pro.docx'):
        """Tạo báo cáo hoàn chỉnh"""
        print("\n" + "="*60)
        print("🚀 BẮT ĐẦU TẠO BÁO CÁO PRO")
        print("="*60 + "\n")
        
        self.add_cover_page()
        self.add_acknowledgment()
        self.add_commitment()
        self.add_table_of_contents()
        self.add_introduction()
        
        # Chương 2: Thuật toán
        print("\n📚 CHƯƠNG 2: Thuật toán tìm kiếm...")
        self.doc.add_heading('CHƯƠNG 2', 1)
        self.doc.add_heading('CÁC THUẬT TOÁN TÌM KIẾM', 1)
        
        search_algos = ['DFS', 'BFS', 'A*', 'Dijkstra']
        for i, algo in enumerate(search_algos, 1):
            if algo in self.config.get('algorithms', []):
                self.add_algorithm_chapter(algo, 2, i)
        
        self.doc.add_page_break()
        
        # Chương 3: CSP
        if 'CSP' in self.config.get('algorithms', []):
            print("\n📚 CHƯƠNG 3: Bài toán ràng buộc...")
            self.doc.add_heading('CHƯƠNG 3', 1)
            self.doc.add_heading('GIẢI BÀI TOÁN RÀNG BUỘC', 1)
            self.add_algorithm_chapter('CSP', 3, 1)
            self.doc.add_page_break()
        
        self.add_conclusion()
        self.add_references()
        
        # Lưu file
        print(f"\n💾 Đang lưu file: {output_file}...")
        self.doc.save(output_file)
        
        print("\n" + "="*60)
        print("✅ HOÀN THÀNH!")
        print("="*60)
        print(f"\n📄 File đã tạo: {output_file}")
        print(f"📊 Biểu đồ: diagrams/")
        print(f"\n💡 Mở file bằng Microsoft Word để xem\n")


def main():
    parser = argparse.ArgumentParser(description='Pro Report Generator')
    parser.add_argument('--mode', choices=['full', 'pdf', 'diagrams'], 
                       default='full', help='Chế độ tạo báo cáo')
    parser.add_argument('--config', default='report_config.json',
                       help='File cấu hình')
    
    args = parser.parse_args()
    
    print("\n" + "╔" + "="*58 + "╗")
    print("║" + " "*58 + "║")
    print("║" + "  🚀 CÔNG CỤ TỰ ĐỘNG TẠO BÁO CÁO PRO VERSION 🚀  ".center(58) + "║")
    print("║" + " "*58 + "║")
    print("╚" + "="*58 + "╝\n")
    
    try:
        generator = ProReportGenerator(args.config)
        
        if args.mode == 'full':
            generator.generate_full_report()
        elif args.mode == 'diagrams':
            print("📊 Chỉ tạo biểu đồ...")
            # Test diagrams
            from diagram_generator import DiagramGenerator
            dg = DiagramGenerator()
            diagrams = [
                ('maze', dg.create_maze_example()),
                ('bfs', dg.create_bfs_visualization()),
                ('astar', dg.create_astar_comparison())
            ]
            for name, buf in diagrams:
                with open(f'diagrams/test_{name}.png', 'wb') as f:
                    f.write(buf.read())
            print("✅ Đã tạo biểu đồ test!")
        
    except KeyboardInterrupt:
        print("\n\n⚠️  Đã hủy bởi người dùng")
    except Exception as e:
        print(f"\n❌ Lỗi: {e}")
        import traceback
        traceback.print_exc()


if __name__ == '__main__':
    main()
