"""
Document Builder
Low-level document construction utilities
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table
from typing import List, Optional, Dict, Any
import os


class DocumentBuilder:
    """Build document elements with consistent styling"""
    
    def __init__(self, doc: Document, style_manager):
        """
        Initialize document builder
        
        Args:
            doc: Document instance
            style_manager: StyleManager instance
        """
        self.doc = doc
        self.style_manager = style_manager
    
    def add_heading(self, text: str, level: int = 1, numbered: bool = False,
                   number: str = None) -> Any:
        """
        Add heading with optional numbering
        
        Args:
            text: Heading text
            level: Heading level (1-3)
            numbered: Whether to add numbering
            number: Custom number (e.g., "1.1", "2.3.1")
            
        Returns:
            Paragraph object
        """
        if numbered and number:
            text = f"{number}. {text}"
        
        return self.doc.add_heading(text, level)
    
    def add_paragraph(self, text: str, style: str = 'Normal',
                     alignment: Optional[Any] = None,
                     bold: bool = False, italic: bool = False) -> Any:
        """
        Add paragraph with styling
        
        Args:
            text: Paragraph text
            style: Paragraph style name
            alignment: Text alignment
            bold: Bold text
            italic: Italic text
            
        Returns:
            Paragraph object
        """
        p = self.doc.add_paragraph(text, style=style)
        
        if alignment:
            p.alignment = alignment
        
        if bold or italic:
            for run in p.runs:
                run.bold = bold
                run.italic = italic
        
        return p
    
    def add_bullet_list(self, items: List[str]) -> None:
        """
        Add bullet list
        
        Args:
            items: List items
        """
        for item in items:
            self.doc.add_paragraph(item, style='List Bullet')
    
    def add_numbered_list(self, items: List[str]) -> None:
        """
        Add numbered list
        
        Args:
            items: List items
        """
        for item in items:
            self.doc.add_paragraph(item, style='List Number')
    
    def add_image(self, image_path: str, width: Optional[Inches] = None,
                 caption: Optional[str] = None) -> None:
        """
        Add image with optional caption
        
        Args:
            image_path: Path to image file
            width: Image width (default: 5.5 inches)
            caption: Image caption
        """
        if not os.path.exists(image_path):
            return
        
        width = width or Inches(5.5)
        
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=width)
        
        if caption:
            self.add_caption(caption)
    
    def add_caption(self, text: str, prefix: str = "Figure") -> Any:
        """
        Add caption
        
        Args:
            text: Caption text
            prefix: Caption prefix (Figure, Table, etc.)
            
        Returns:
            Paragraph object
        """
        caption_style = self.style_manager.get_style('caption')
        
        p = self.doc.add_paragraph()
        p.alignment = caption_style.get('alignment', WD_ALIGN_PARAGRAPH.CENTER)
        
        run = p.add_run(f"{prefix}: {text}")
        run.font.size = caption_style.get('font_size', Pt(11))
        run.italic = caption_style.get('italic', True)
        
        return p
    
    def add_table(self, headers: List[str], rows: List[List[Any]],
                 style: str = 'Light Grid Accent 1') -> Table:
        """
        Add table
        
        Args:
            headers: Column headers
            rows: Table rows
            style: Table style name
            
        Returns:
            Table object
        """
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = style
        
        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Bold header text
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add rows
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
        
        return table
    
    def add_code_block(self, code: str) -> Any:
        """
        Add code block with styling
        
        Args:
            code: Code text
            
        Returns:
            Paragraph object
        """
        code_style = self.style_manager.get_style('code')
        
        p = self.doc.add_paragraph()
        run = p.add_run(code)
        run.font.name = code_style.get('font_name', 'Courier New')
        run.font.size = code_style.get('font_size', Pt(10))
        
        # Add background shading
        if 'background_color' in code_style:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), code_style['background_color'])
            p._element.get_or_add_pPr().append(shading)
        
        # Add indentation
        if 'left_indent' in code_style:
            p.paragraph_format.left_indent = code_style['left_indent']
        
        return p
    
    def add_page_break(self) -> None:
        """Add page break"""
        self.doc.add_page_break()
    
    def add_toc(self) -> None:
        """
        Add Table of Contents field
        
        Note: TOC will be generated when document is opened in Word
        """
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        
        # Create TOC field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        r_element = run._element
        r_element.append(fldChar1)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar3)
        
        # Add instruction text
        p = self.doc.add_paragraph()
        p.add_run("Right-click and select 'Update Field' to generate TOC").italic = True
    
    def add_centered_text(self, text: str, size: int = 13,
                         bold: bool = False) -> Any:
        """
        Add centered text
        
        Args:
            text: Text content
            size: Font size
            bold: Bold text
            
        Returns:
            Paragraph object
        """
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        return p
