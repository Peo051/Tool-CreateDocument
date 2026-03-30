import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))

from models import Report, Section, SectionType, EvidenceType

class DocxRenderer:
    """Render report to DOCX format"""
    
    def __init__(self):
        self.doc = None
    
    def render(self, report: Report, output_path: str) -> str:
        """Render report to DOCX file"""
        self.doc = Document()
        self._setup_document()
        
        # Render sections
        for section in report.sections:
            self._render_section(section, report)
        
        # Save
        os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
        self.doc.save(output_path)
        return output_path
    
    def _setup_document(self):
        """Setup document formatting"""
        sections = self.doc.sections
        for section in sections:
            section.page_height = Cm(29.7)
            section.page_width = Cm(21)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2)
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
        
        self._setup_styles()
    
    def _setup_styles(self):
        """Setup text styles"""
        styles = self.doc.styles
        
        normal = styles['Normal']
        normal.font.name = 'Times New Roman'
        normal.font.size = Pt(13)
        normal.paragraph_format.line_spacing = 1.5
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        for i in range(1, 4):
            h = styles[f'Heading {i}']
            h.font.name = 'Times New Roman'
            h.font.bold = True
            if i == 1:
                h.font.size = Pt(16)
                h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif i == 2:
                h.font.size = Pt(14)
            else:
                h.font.size = Pt(13)
    
    def _render_section(self, section: Section, report: Report):
        """Render a section"""
        
        if section.type == SectionType.COVER:
            self._render_cover(section, report)
        elif section.type == SectionType.TOC:
            self._render_toc(section)
        elif section.type == SectionType.ACKNOWLEDGMENT:
            self._render_standard_section(section)
        elif section.type == SectionType.COMMITMENT:
            self._render_commitment(section, report)
        elif section.type == SectionType.INTRODUCTION:
            self._render_introduction(section)
        elif section.type == SectionType.CHAPTER:
            self._render_chapter(section, report)
        elif section.type == SectionType.ALGORITHM:
            self._render_algorithm(section, report)
        elif section.type == SectionType.CONCLUSION:
            self._render_standard_section(section)
        elif section.type == SectionType.REFERENCES:
            self._render_references(section)
        else:
            self._render_standard_section(section)
    
    def _render_cover(self, section: Section, report: Report):
        """Render cover page"""
        meta = report.metadata
        config = section.metadata
        
        self._add_centered_text('BỘ CÔNG THƯƠNG', 13, True)
        self._add_centered_text(f'TRƯỜNG {meta.organization.upper()}', 13, True)
        self._add_centered_text(f'KHOA {meta.faculty.upper()}', 13, True)
        self._add_centered_text('─' * 30, 13)
        
        for _ in range(3):
            self.doc.add_paragraph()
        
        self._add_centered_text('BÁO CÁO ĐỒ ÁN MÔN HỌC', 16, True)
        self._add_centered_text('TRÍ TUỆ NHÂN TẠO', 16, True)
        self.doc.add_paragraph()
        
        self._add_centered_text(f'ĐỀ TÀI: {meta.title}', 14, True)
        if meta.subtitle:
            self._add_centered_text(meta.subtitle, 13)
        
        for _ in range(3):
            self.doc.add_paragraph()
        
        if meta.advisor:
            self._add_centered_text(f'Giảng viên hướng dẫn: {meta.advisor}', 13)
        self.doc.add_paragraph()
        
        self._add_centered_text('Sinh viên thực hiện:', 13, True)
        for author in meta.authors:
            name = author.get('name', '')
            student_id = author.get('id', '')
            self._add_centered_text(f'{name} - MSSV: {student_id}', 13)
        
        for _ in range(2):
            self.doc.add_paragraph()
        
        if meta.class_name:
            self._add_centered_text(f'Lớp: {meta.class_name}', 13)
        
        for _ in range(2):
            self.doc.add_paragraph()
        
        now = datetime.now()
        self._add_centered_text(f'TP. HỒ CHÍ MINH, tháng {now.month} năm {now.year}', 13)
        
        self.doc.add_page_break()
    
    def _render_toc(self, section: Section):
        """Render table of contents"""
        self.doc.add_heading(section.title, 1)
        
        toc_items = [
            "PHẦN MỞ ĐẦU",
            "CHƯƠNG 1: TỔNG QUAN",
            "CHƯƠNG 2: THUẬT TOÁN TÌM KIẾM",
            "CHƯƠNG 3: BÀI TOÁN RÀNG BUỘC",
            "KẾT LUẬN",
            "TÀI LIỆU THAM KHẢO"
        ]
        
        for item in toc_items:
            p = self.doc.add_paragraph()
            run = p.add_run(item)
            run.bold = True
        
        self.doc.add_page_break()
    
    def _render_commitment(self, section: Section, report: Report):
        """Render commitment section"""
        self.doc.add_heading(section.title, section.level)
        self.doc.add_paragraph(section.content)
        self.doc.add_paragraph()
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run('Sinh viên thực hiện')
        run.bold = True
        
        for author in report.metadata.authors:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(author.get('name', ''))
        
        self.doc.add_page_break()
    
    def _render_standard_section(self, section: Section):
        """Render standard section with content"""
        self.doc.add_heading(section.title, section.level)
        if section.content:
            self.doc.add_paragraph(section.content)
        self.doc.add_page_break()
    
    def _render_introduction(self, section: Section):
        """Render introduction with subsections"""
        self.doc.add_heading(section.title, section.level)
        
        for subsection in section.subsections:
            self.doc.add_heading(subsection.title, subsection.level)
            if subsection.content:
                # Handle bullet points
                if '\n-' in subsection.content:
                    lines = subsection.content.split('\n')
                    for line in lines:
                        if line.strip().startswith('-'):
                            self.doc.add_paragraph(line.strip()[1:].strip(), style='List Bullet')
                        elif line.strip():
                            self.doc.add_paragraph(line.strip())
                else:
                    self.doc.add_paragraph(subsection.content)
        
        self.doc.add_page_break()
    
    def _render_chapter(self, section: Section, report: Report):
        """Render chapter with subsections"""
        self.doc.add_heading(section.title, section.level)
        
        for subsection in section.subsections:
            self._render_section(subsection, report)
        
        self.doc.add_page_break()
    
    def _render_algorithm(self, section: Section, report: Report):
        """Render algorithm section"""
        self.doc.add_heading(section.title, section.level)
        
        if section.content:
            # Parse and render formatted content
            lines = section.content.split('\n')
            in_code_block = False
            code_lines = []
            
            for line in lines:
                if line.strip() == '```':
                    if in_code_block:
                        # End code block
                        self._add_code_block('\n'.join(code_lines))
                        code_lines = []
                        in_code_block = False
                    else:
                        # Start code block
                        in_code_block = True
                elif in_code_block:
                    code_lines.append(line)
                elif line.strip().startswith('**') and line.strip().endswith('**'):
                    # Bold heading
                    text = line.strip()[2:-2]
                    p = self.doc.add_paragraph()
                    run = p.add_run(text)
                    run.bold = True
                elif line.strip().startswith('-'):
                    # Bullet point
                    self.doc.add_paragraph(line.strip()[1:].strip(), style='List Bullet')
                elif line.strip():
                    self.doc.add_paragraph(line.strip())
        
        # Add diagram if exists
        self._add_evidence_for_section(section, report)
    
    def _render_references(self, section: Section):
        """Render references"""
        self.doc.add_heading(section.title, section.level or 1)
        
        if section.content:
            # Content is already formatted by bibliography generator
            lines = section.content.split('\n\n')
            for line in lines:
                if line.strip():
                    self.doc.add_paragraph(line.strip())
        elif hasattr(section, 'metadata') and 'references' in section.metadata:
            # Fallback: render from metadata
            for ref_data in section.metadata.get('references', []):
                citation = ref_data.get('citation', '')
                if citation:
                    self.doc.add_paragraph(citation)
    
    def _add_evidence_for_section(self, section: Section, report: Report):
        """Add evidence (diagrams) for section"""
        for evidence in report.evidence:
            if evidence.position == section.id and evidence.type == EvidenceType.FIGURE:
                if os.path.exists(evidence.file_path):
                    self.doc.add_paragraph()
                    self.doc.add_picture(evidence.file_path, width=Inches(5.5))
                    
                    if evidence.caption:
                        p = self.doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run(evidence.caption)
                        run.italic = True
                        run.font.size = Pt(11)
    
    def _add_centered_text(self, text: str, size: int, bold: bool = False):
        """Add centered text"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
    
    def _add_code_block(self, code: str):
        """Add code block"""
        p = self.doc.add_paragraph()
        run = p.add_run(code)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F0F0F0')
        p._element.get_or_add_pPr().append(shading)
        
        p.paragraph_format.left_indent = Inches(0.5)
