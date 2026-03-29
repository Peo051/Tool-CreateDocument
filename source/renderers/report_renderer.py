"""
Report Renderer
High-level report rendering to DOCX format
"""

import os
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Optional, Dict, Any

from .style_manager import StyleManager
from .document_builder import DocumentBuilder


class ReportRenderer:
    """Render structured report to DOCX document"""
    
    def __init__(self, style_manager: Optional[StyleManager] = None):
        """
        Initialize renderer
        
        Args:
            style_manager: Custom style manager (optional)
        """
        self.style_manager = style_manager or StyleManager()
        self.doc = None
        self.builder = None
        self.section_counters = {}
    
    def render(self, report_data: Dict[str, Any], output_path: str) -> str:
        """
        Render report to DOCX file
        
        Args:
            report_data: Report data dictionary with structure:
                - metadata: Project metadata
                - sections: List of sections
                - evidence: List of evidence items
            output_path: Output file path
            
        Returns:
            Output file path
        """
        # Create document
        self.doc = Document()
        self.style_manager.apply_to_document(self.doc)
        self.builder = DocumentBuilder(self.doc, self.style_manager)
        
        # Reset counters
        self.section_counters = {'chapter': 0, 'section': 0, 'subsection': 0}
        
        # Render sections
        metadata = report_data.get('metadata', {})
        sections = report_data.get('sections', [])
        evidence = report_data.get('evidence', {})
        
        for section in sections:
            self._render_section(section, metadata, evidence)
        
        # Save document
        os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
        self.doc.save(output_path)
        
        return output_path
    
    def _render_section(self, section: Dict[str, Any], metadata: Dict[str, Any],
                       evidence: Dict[str, Any]) -> None:
        """Render a section"""
        section_type = section.get('type', 'standard')
        
        renderers = {
            'cover': self._render_cover,
            'toc': self._render_toc,
            'acknowledgment': self._render_standard,
            'commitment': self._render_commitment,
            'chapter': self._render_chapter,
            'section': self._render_standard,
            'conclusion': self._render_standard,
            'references': self._render_references,
        }
        
        renderer = renderers.get(section_type, self._render_standard)
        renderer(section, metadata, evidence)
    
    def _render_cover(self, section: Dict[str, Any], metadata: Dict[str, Any],
                     evidence: Dict[str, Any]) -> None:
        """Render cover page"""
        # Organization header
        org = metadata.get('organization', metadata.get('university', 'UNIVERSITY'))
        faculty = metadata.get('faculty', 'FACULTY')
        
        self.builder.add_centered_text('BỘ CÔNG THƯƠNG', 13, True)
        self.builder.add_centered_text(f'TRƯỜNG {org.upper()}', 13, True)
        self.builder.add_centered_text(f'KHOA {faculty.upper()}', 13, True)
        self.builder.add_centered_text('─' * 30, 13)
        
        # Spacing
        for _ in range(3):
            self.doc.add_paragraph()
        
        # Report type
        report_type = metadata.get('report_type', 'project_report')
        if report_type == 'thesis':
            self.builder.add_centered_text('LUẬN VĂN TỐT NGHIỆP', 16, True)
        elif report_type == 'technical_report':
            self.builder.add_centered_text('BÁO CÁO KỸ THUẬT', 16, True)
        else:
            self.builder.add_centered_text('BÁO CÁO ĐỒ ÁN', 16, True)
        
        self.doc.add_paragraph()
        
        # Title
        title = metadata.get('title', 'PROJECT TITLE')
        self.builder.add_centered_text(f'ĐỀ TÀI: {title}', 14, True)
        
        subtitle = metadata.get('subtitle')
        if subtitle:
            self.builder.add_centered_text(subtitle, 13)
        
        # Spacing
        for _ in range(3):
            self.doc.add_paragraph()
        
        # Advisor
        advisor = metadata.get('advisor')
        if advisor:
            self.builder.add_centered_text(f'Giảng viên hướng dẫn: {advisor}', 13)
        
        self.doc.add_paragraph()
        
        # Authors
        self.builder.add_centered_text('Sinh viên thực hiện:', 13, True)
        authors = metadata.get('students', metadata.get('authors', []))
        for author in authors:
            if isinstance(author, dict):
                name = author.get('name', '')
                student_id = author.get('id', '')
                self.builder.add_centered_text(f'{name} - MSSV: {student_id}', 13)
            else:
                self.builder.add_centered_text(str(author), 13)
        
        # Spacing
        for _ in range(2):
            self.doc.add_paragraph()
        
        # Class
        class_name = metadata.get('class_name', metadata.get('class'))
        if class_name:
            self.builder.add_centered_text(f'Lớp: {class_name}', 13)
        
        # Spacing
        for _ in range(2):
            self.doc.add_paragraph()
        
        # Date
        now = datetime.now()
        self.builder.add_centered_text(
            f'TP. HỒ CHÍ MINH, tháng {now.month} năm {now.year}', 13
        )
        
        self.builder.add_page_break()
    
    def _render_toc(self, section: Dict[str, Any], metadata: Dict[str, Any],
                   evidence: Dict[str, Any]) -> None:
        """Render table of contents"""
        title = section.get('title', 'MỤC LỤC')
        self.builder.add_heading(title, 1)
        
        # Add TOC field
        self.builder.add_toc()
        
        self.builder.add_page_break()
    
    def _render_commitment(self, section: Dict[str, Any], metadata: Dict[str, Any],
                          evidence: Dict[str, Any]) -> None:
        """Render commitment section"""
        title = section.get('title', 'LỜI CAM ĐOAN')
        content = section.get('content', '')
        
        self.builder.add_heading(title, 1)
        
        if content:
            self.builder.add_paragraph(content)
        
        self.doc.add_paragraph()
        
        # Signature
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run('Sinh viên thực hiện')
        run.bold = True
        
        authors = metadata.get('students', metadata.get('authors', []))
        for author in authors:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if isinstance(author, dict):
                p.add_run(author.get('name', ''))
            else:
                p.add_run(str(author))
        
        self.builder.add_page_break()
    
    def _render_chapter(self, section: Dict[str, Any], metadata: Dict[str, Any],
                       evidence: Dict[str, Any]) -> None:
        """Render chapter with subsections"""
        title = section.get('title', '')
        content = section.get('content', '')
        subsections = section.get('subsections', [])
        level = section.get('level', 1)
        
        # Increment chapter counter
        self.section_counters['chapter'] += 1
        chapter_num = self.section_counters['chapter']
        
        # Add chapter heading
        self.builder.add_heading(title, level, numbered=True, number=f"CHƯƠNG {chapter_num}")
        
        # Add chapter content
        if content:
            self._render_content(content)
        
        # Reset section counter for new chapter
        self.section_counters['section'] = 0
        
        # Render subsections
        for subsection in subsections:
            self._render_subsection(subsection, evidence, chapter_num)
        
        self.builder.add_page_break()
    
    def _render_subsection(self, section: Dict[str, Any], evidence: Dict[str, Any],
                          chapter_num: int) -> None:
        """Render subsection"""
        title = section.get('title', '')
        content = section.get('content', '')
        level = section.get('level', 2)
        section_id = section.get('id', '')
        
        # Increment section counter
        self.section_counters['section'] += 1
        section_num = self.section_counters['section']
        
        # Add section heading
        number = f"{chapter_num}.{section_num}"
        self.builder.add_heading(title, level, numbered=True, number=number)
        
        # Add content
        if content:
            self._render_content(content)
        
        # Add evidence for this section
        self._add_evidence_for_section(section_id, evidence)
    
    def _render_standard(self, section: Dict[str, Any], metadata: Dict[str, Any],
                        evidence: Dict[str, Any]) -> None:
        """Render standard section"""
        title = section.get('title', '')
        content = section.get('content', '')
        level = section.get('level', 1)
        subsections = section.get('subsections', [])
        
        # Add heading
        self.builder.add_heading(title, level)
        
        # Add content
        if content:
            self._render_content(content)
        
        # Render subsections
        for subsection in subsections:
            sub_title = subsection.get('title', '')
            sub_content = subsection.get('content', '')
            sub_level = subsection.get('level', level + 1)
            
            if sub_title:
                self.builder.add_heading(sub_title, sub_level)
            
            if sub_content:
                self._render_content(sub_content)
        
        self.builder.add_page_break()
    
    def _render_references(self, section: Dict[str, Any], metadata: Dict[str, Any],
                          evidence: Dict[str, Any]) -> None:
        """Render references section"""
        title = section.get('title', 'TÀI LIỆU THAM KHẢO')
        content = section.get('content', '')
        references = section.get('references', [])
        
        self.builder.add_heading(title, 1)
        
        if references:
            self.builder.add_numbered_list(references)
        elif content:
            # Parse references from content
            refs = [r.strip() for r in content.split('\n') if r.strip()]
            self.builder.add_numbered_list(refs)
        
        self.builder.add_page_break()
    
    def _render_content(self, content: str) -> None:
        """Render content with formatting"""
        if not content:
            return
        
        lines = content.split('\n')
        in_code_block = False
        code_lines = []
        
        for line in lines:
            stripped = line.strip()
            
            # Code block markers
            if stripped == '```' or stripped.startswith('```'):
                if in_code_block:
                    # End code block
                    self.builder.add_code_block('\n'.join(code_lines))
                    code_lines = []
                    in_code_block = False
                else:
                    # Start code block
                    in_code_block = True
                continue
            
            if in_code_block:
                code_lines.append(line)
                continue
            
            # Bullet points
            if stripped.startswith('- ') or stripped.startswith('* '):
                self.doc.add_paragraph(stripped[2:], style='List Bullet')
            # Numbered lists
            elif stripped and stripped[0].isdigit() and '. ' in stripped:
                self.doc.add_paragraph(stripped.split('. ', 1)[1], style='List Number')
            # Regular paragraphs
            elif stripped:
                self.builder.add_paragraph(stripped)
    
    def _add_evidence_for_section(self, section_id: str, evidence: Dict[str, Any]) -> None:
        """Add evidence (images, tables) for section"""
        # Images
        figures = evidence.get('figures', [])
        for figure in figures:
            if figure.get('section_id') == section_id or figure.get('position') == section_id:
                image_path = figure.get('path', figure.get('file_path'))
                caption = figure.get('caption', figure.get('title'))
                
                if image_path:
                    self.builder.add_image(image_path, caption=caption)
        
        # Tables
        tables = evidence.get('tables', [])
        for table in tables:
            if table.get('section_id') == section_id or table.get('position') == section_id:
                headers = table.get('headers', [])
                rows = table.get('rows', [])
                caption = table.get('caption', table.get('title'))
                
                if headers and rows:
                    self.builder.add_table(headers, rows)
                    if caption:
                        self.builder.add_caption(caption, prefix="Table")
